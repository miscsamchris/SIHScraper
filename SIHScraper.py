from bs4 import BeautifulSoup
import requests
import re
from docx import Document
from docx.shared import Inches
table=[]
f=requests.get("https://www.sih.gov.in/sih2019ProblemStatements")
soup = BeautifulSoup(f.content, 'html.parser')
r=soup.find("tbody")
x=r.findChildren("tr",recursive=False)
for i in x:
    j=i.find_all("td")
    row=[]
    for y in j:
        bs=BeautifulSoup(str(y),"html.parser")
        if j.index(y) in [2]:
            a=bs.get_text().split("\n",1)[0]
            row.append(a)
        if j.index(y) in [1,3,4,7,8,10,11]:
            a=bs.get_text()
            a=re.sub(r"\n{2,20}","\n",a)
            a=re.sub(r" View","",a)
            a=re.sub(r"Problem Statement Details","",a)
            a=a.replace("×","")
            a=re.sub(r"\n{2,20}","\n",a)
            row.append(a)
    table.append(row)
for i in range(2,51):
    f=requests.get("https://www.sih.gov.in/sih2019ProblemStatements?page="+str(i))
    soup = BeautifulSoup(f.content, 'html.parser')
    r=soup.find("tbody")
    x=r.findChildren("tr",recursive=False)
    for i in x:
        j=i.find_all("td")
        row=[]
        for y in j:
            bs=BeautifulSoup(str(y),"html.parser")
            if j.index(y) in [2]:
                a=bs.get_text().split("\n",1)[0]
                row.append(a)
            if j.index(y) in [1,3,4,7,8,10,11]:
                a=bs.get_text()
                a=re.sub(r"\n{2,20}","\n",a)
                a=re.sub(r" View","",a)
                a=re.sub(r"Problem Statement Details","",a)
                a=a.replace("×","")
                a=re.sub(r"\n{2,20}","\n",a)
                row.append(a)
        table.append(row)
document = Document()
document.add_heading('Smart India Hackathon 2019 Problem Statements', 0)

for i in table:
        document.add_heading(str(table.index(i)+1)+".  "+i[1],level=1)
        for j in i:
            if i.index(j)==0:
                document.add_paragraph("Organisation : "+str(j))
            if i.index(j)==1:
                document.add_paragraph("Title : "+str(j))        
            if i.index(j)==5:
                document.add_paragraph("Category : "+str(j))
            if i.index(j)==6:
                document.add_paragraph("Technology : "+str(j))
            if i.index(j)==3:
                document.add_paragraph("Organisation Category : "+str(j).strip())
            if i.index(j)==7:
                document.add_paragraph("Complexity : "+str(j))
            if i.index(j)==4:
                p=document.add_paragraph("Youtube link : ")
                if len(str(j))>13:
                     p.add_run(str(j))
            if i.index(j)==2:
                document.add_heading("Problem Statement Details",level=4)
                document.add_paragraph(str(j))
document.save('SIHProblemStatements.docx')
